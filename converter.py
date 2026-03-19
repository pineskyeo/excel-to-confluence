"""
Excel → Word (DOCX) Converter for Confluence Import
Usage: python converter.py <input.xlsx> [output.docx]
"""
import io
import os
import sys

import openpyxl
from openpyxl.drawing.spreadsheet_drawing import OneCellAnchor, TwoCellAnchor
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage


# ── helpers ────────────────────────────────────────────────────────────────────

def _hex_to_rgb(hex_str: str):
    """'FFAABBCC' or 'AABBCC' → RGBColor"""
    h = hex_str.lstrip("#")
    if len(h) == 8:          # ARGB
        h = h[2:]
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return RGBColor(r, g, b)


def _set_cell_shading(doc_cell, hex_color: str):
    """Apply background fill to a Word table cell."""
    h = hex_color.lstrip("#")
    if len(h) == 8:
        h = h[2:]
    tc = doc_cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove existing shd
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), h.upper())
    tcPr.append(shd)


def _image_anchor_row(img) -> int:
    """Return 1-based row of an openpyxl image anchor (best effort)."""
    anchor = img.anchor
    if isinstance(anchor, (OneCellAnchor, TwoCellAnchor)):
        return anchor._from.row + 1          # _from.row is 0-based
    return 1                                 # AbsoluteAnchor fallback


def _load_image_bytes(img):
    """Extract raw bytes from an openpyxl Image object."""
    try:
        data = img._data()
        if callable(data):
            data = data()
        return bytes(data)
    except Exception as exc:
        print(f"  [warn] could not read image data: {exc}")
        return None


def _to_png_bytes(raw: bytes):
    """Ensure image bytes are in a format python-docx can handle."""
    try:
        buf_in = io.BytesIO(raw)
        pil = PILImage.open(buf_in)
        pil.load()
        out = io.BytesIO()
        fmt = pil.format if pil.format in ("PNG", "JPEG", "BMP") else "PNG"
        if pil.mode not in ("RGB", "RGBA", "L"):
            pil = pil.convert("RGB")
        pil.save(out, format=fmt)
        out.seek(0)
        return out
    except Exception as exc:
        print(f"  [warn] image conversion failed: {exc}")
        return None


# ── content-block detection ────────────────────────────────────────────────────

class Block:
    """A positioned content block on a sheet."""
    __slots__ = ("kind", "row", "payload")

    def __init__(self, kind: str, row: int, payload):
        self.kind = kind      # "heading" | "text" | "table" | "image"
        self.row = row
        self.payload = payload


def _scan_sheet(ws) -> list[Block]:
    """
    Returns an ordered list of Blocks derived from the worksheet.
    Strategy:
      1. Named Excel tables  → 'table' blocks
      2. Images              → 'image' blocks
      3. Remaining non-empty rows that aren't inside a table:
         - single-column or looks like prose → 'text' / 'heading'
         - multi-column → small inline table
    """
    blocks: list[Block] = []

    # ── 1. Named tables ────────────────────────────────────────────────
    table_row_sets: list[set] = []
    for tbl in ws.tables.values():
        ref = ws[tbl.ref]          # list of tuples of cells
        min_r = ref[0][0].row
        max_r = ref[-1][0].row
        min_c = ref[0][0].column
        max_c = ref[0][-1].column
        blocks.append(Block("table", min_r, {
            "min_row": min_r, "max_row": max_r,
            "min_col": min_c, "max_col": max_c,
        }))
        table_row_sets.append(set(range(min_r, max_r + 1)))

    # ── 2. Images ──────────────────────────────────────────────────────
    for img in ws._images:
        raw = _load_image_bytes(img)
        if raw is None:
            continue
        row = _image_anchor_row(img)
        pixel_w = getattr(img, "width", 300) or 300
        blocks.append(Block("image", row, {"raw": raw, "pixel_w": pixel_w}))

    # ── 3. Text / auto-detected tables from remaining rows ─────────────
    in_named_table = set()
    for s in table_row_sets:
        in_named_table |= s

    # group consecutive data rows not already claimed
    pending: list[tuple] = []      # (row_num, [cells_with_values])
    for ws_row in ws.iter_rows():
        r = ws_row[0].row
        if r in in_named_table:
            if pending:
                _flush_pending(pending, blocks)
                pending = []
            continue
        filled = [c for c in ws_row if c.value not in (None, "")]
        if filled:
            pending.append((r, filled))
        else:
            if pending:
                _flush_pending(pending, blocks)
                pending = []
    if pending:
        _flush_pending(pending, blocks)

    blocks.sort(key=lambda b: b.row)
    return blocks


def _flush_pending(pending: list, blocks: list):
    """
    Decide whether pending rows form a table or prose text and emit blocks.
    Heuristic: if most rows span >= 2 columns → table; else text/heading.
    """
    if not pending:
        return

    multi_col = sum(1 for _, cells in pending if len(cells) >= 2)
    is_table = multi_col / len(pending) >= 0.5 and len(pending) >= 2

    first_row = pending[0][0]

    if is_table:
        all_cols = [c.column for _, cells in pending for c in cells]
        blocks.append(Block("table", first_row, {
            "rows": pending,          # raw rows
            "min_row": first_row,
            "max_row": pending[-1][0],
            "min_col": min(all_cols),
            "max_col": max(all_cols),
            "raw_mode": True,         # signal to renderer
        }))
    else:
        for row_num, cells in pending:
            cell = cells[0]
            val = str(cell.value)
            is_bold = cell.font and cell.font.bold
            size = (cell.font.size or 11) if cell.font else 11
            kind = "heading" if (is_bold and size >= 13) else "text"
            blocks.append(Block(kind, row_num, {
                "text": val,
                "bold": is_bold,
                "italic": cell.font.italic if cell.font else False,
                "size": size,
            }))


# ── Word document builder ──────────────────────────────────────────────────────

class WordBuilder:
    def __init__(self):
        self.doc = Document()
        # page margins (narrower to fit wide tables)
        for section in self.doc.sections:
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

    # ── public entry points ────────────────────────────────────────────

    def add_sheet_heading(self, title: str):
        self.doc.add_heading(title, level=1)

    def add_block(self, block: Block, ws):
        if block.kind == "heading":
            self._add_heading(block.payload)
        elif block.kind == "text":
            self._add_paragraph(block.payload)
        elif block.kind == "table":
            self._add_table(block.payload, ws)
        elif block.kind == "image":
            self._add_image(block.payload)

    def save(self, path: str):
        self.doc.save(path)

    # ── private renderers ──────────────────────────────────────────────

    def _add_heading(self, p: dict):
        size = p.get("size", 11)
        level = 2 if size >= 16 else (3 if size >= 13 else 4)
        self.doc.add_heading(p["text"], level=level)

    def _add_paragraph(self, p: dict):
        para = self.doc.add_paragraph()
        run = para.add_run(p["text"])
        run.bold = p.get("bold", False)
        run.italic = p.get("italic", False)
        if p.get("size"):
            run.font.size = Pt(p["size"])

    def _add_table(self, p: dict, ws):
        if p.get("raw_mode"):
            self._add_table_raw(p)
        else:
            self._add_table_ws(p, ws)
        # spacer
        self.doc.add_paragraph()

    def _add_table_ws(self, p: dict, ws):
        """Render a named Excel table from worksheet coordinates."""
        min_r, max_r = p["min_row"], p["max_row"]
        min_c, max_c = p["min_col"], p["max_col"]
        n_rows = max_r - min_r + 1
        n_cols = max_c - min_c + 1

        # Build merged-cell map  { (row, col) → MergedCell range }
        merged_map = {}
        for mr in ws.merged_cells.ranges:
            for r in range(mr.min_row, mr.max_row + 1):
                for c in range(mr.min_col, mr.max_col + 1):
                    merged_map[(r, c)] = mr

        tbl = self.doc.add_table(rows=n_rows, cols=n_cols)
        tbl.style = "Table Grid"

        for i, row_num in enumerate(range(min_r, max_r + 1)):
            for j, col_num in enumerate(range(min_c, max_c + 1)):
                cell = ws.cell(row=row_num, column=col_num)
                doc_cell = tbl.cell(i, j)

                val = cell.value if cell.value is not None else ""
                para = doc_cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(str(val))

                # font
                if cell.font:
                    run.bold = bool(cell.font.bold)
                    run.italic = bool(cell.font.italic)
                    if cell.font.color and cell.font.color.type == "rgb":
                        try:
                            run.font.color.rgb = _hex_to_rgb(cell.font.color.rgb)
                        except Exception:
                            pass

                # background
                if cell.fill and cell.fill.patternType == "solid":
                    try:
                        bg = cell.fill.fgColor.rgb
                        if bg and bg not in ("00000000", "FF000000", "000000"):
                            _set_cell_shading(doc_cell, bg)
                    except Exception:
                        pass

        # handle merges within the range
        for mr in ws.merged_cells.ranges:
            # only process merges that are fully inside our table range
            if (mr.min_row >= min_r and mr.max_row <= max_r and
                    mr.min_col >= min_c and mr.max_col <= max_c):
                r1 = mr.min_row - min_r
                c1 = mr.min_col - min_c
                r2 = mr.max_row - min_r
                c2 = mr.max_col - min_c
                try:
                    tbl.cell(r1, c1).merge(tbl.cell(r2, c2))
                except Exception:
                    pass

    def _add_table_raw(self, p: dict):
        """Render an auto-detected table from raw pending rows."""
        rows = p["rows"]
        min_c = p["min_col"]
        max_c = p["max_col"]
        n_cols = max_c - min_c + 1

        tbl = self.doc.add_table(rows=len(rows), cols=n_cols)
        tbl.style = "Table Grid"

        for i, (_, cells) in enumerate(rows):
            # index cells by column
            col_map = {c.column: c for c in cells}
            for j in range(n_cols):
                col_num = min_c + j
                cell = col_map.get(col_num)
                doc_cell = tbl.cell(i, j)
                para = doc_cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if cell is None:
                    para.add_run("")
                    continue

                val = cell.value if cell.value is not None else ""
                run = para.add_run(str(val))

                if cell.font:
                    run.bold = bool(cell.font.bold)
                    run.italic = bool(cell.font.italic)
                    if cell.font.color and cell.font.color.type == "rgb":
                        try:
                            run.font.color.rgb = _hex_to_rgb(cell.font.color.rgb)
                        except Exception:
                            pass

                if cell.fill and cell.fill.patternType == "solid":
                    try:
                        bg = cell.fill.fgColor.rgb
                        if bg and bg not in ("00000000", "FF000000", "000000"):
                            _set_cell_shading(doc_cell, bg)
                    except Exception:
                        pass

    def _add_image(self, p: dict):
        buf = _to_png_bytes(p["raw"])
        if buf is None:
            return
        pixel_w = p.get("pixel_w", 300) or 300
        # cap at 5.5 inches (page width minus margins)
        inch_w = min(pixel_w / 96, 5.5)
        try:
            self.doc.add_picture(buf, width=Inches(inch_w))
        except Exception as exc:
            print(f"  [warn] could not embed image: {exc}")
            return
        # center the picture paragraph
        last_para = self.doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()


# ── main converter ─────────────────────────────────────────────────────────────

class ExcelToWordConverter:
    def __init__(self, xlsx_path: str, docx_path: str):
        self.xlsx_path = xlsx_path
        self.docx_path = docx_path

    def convert(self):
        print(f"📖 Reading: {self.xlsx_path}")
        wb = openpyxl.load_workbook(self.xlsx_path, data_only=True)
        builder = WordBuilder()

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            print(f"  📄 Sheet: {sheet_name}")

            builder.add_sheet_heading(sheet_name)

            blocks = _scan_sheet(ws)
            print(f"     {len(blocks)} block(s) found")

            for block in blocks:
                builder.add_block(block, ws)

        builder.save(self.docx_path)
        print(f"✅ Saved: {self.docx_path}")
        print()
        print("Confluence 업로드 방법:")
        print("  1. Confluence 페이지 편집기 열기")
        print("  2. ··· 메뉴 → 'Word 문서 임포트' 선택")
        print(f"  3. {os.path.basename(self.docx_path)} 업로드")


def main():
    if len(sys.argv) < 2:
        print("Usage: python converter.py <input.xlsx> [output.docx]")
        sys.exit(1)

    xlsx = sys.argv[1]
    docx = sys.argv[2] if len(sys.argv) > 2 else xlsx.replace(".xlsx", ".docx")

    if not os.path.exists(xlsx):
        print(f"Error: file not found — {xlsx}")
        sys.exit(1)

    ExcelToWordConverter(xlsx, docx).convert()


if __name__ == "__main__":
    main()
